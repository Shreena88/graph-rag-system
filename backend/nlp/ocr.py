import numpy as np
import logging
from typing import List
from backend.models.document import PageContent, ParsedDocument

# Suppress pdfminer font warnings
logging.getLogger("pdfminer").setLevel(logging.ERROR)


class DocumentParser:
    def __init__(self, use_gpu: bool = False):
        self._ocr = None  # lazy-loaded
        self._use_gpu = use_gpu

    def _get_ocr(self):
        if self._ocr is None:
            from paddleocr import PaddleOCR
            device = "gpu" if self._use_gpu else "cpu"
            # paddleocr >=2.8 uses device= param; older versions use use_gpu=
            # New API raises ValueError for unknown args, old raises TypeError
            try:
                self._ocr = PaddleOCR(lang="en", device=device)
            except (TypeError, ValueError):
                self._ocr = PaddleOCR(use_angle_cls=True, lang="en", use_gpu=self._use_gpu)
        return self._ocr

    def parse(self, file_path: str, filename: str) -> ParsedDocument:
        if file_path.endswith(".pdf"):
            pages = self._extract_native_pdf(file_path)
            if self._is_scanned(pages):
                pages = self._extract_with_ocr(file_path)
        elif file_path.endswith(".txt"):
            pages = self._extract_text(file_path)
        elif file_path.endswith(".docx"):
            pages = self._extract_docx(file_path)
        else:
            pages = self._extract_with_ocr(file_path)

        return ParsedDocument(pages=pages, filename=filename, total_pages=len(pages))

    def _extract_native_pdf(self, path: str) -> List[PageContent]:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables() or []
                pages.append(PageContent(
                    page_number=i + 1,
                    text=text,
                    tables=[{"rows": t} for t in tables if t],
                ))
        return pages

    def _extract_with_ocr(self, path: str) -> List[PageContent]:
        from pdf2image import convert_from_path
        ocr = self._get_ocr()
        images = convert_from_path(path, dpi=300)
        pages = []
        for i, img in enumerate(images):
            result = ocr.ocr(np.array(img), cls=True)
            text = "\n".join(
                line[1][0] for block in (result or []) for line in (block or [])
            )
            pages.append(PageContent(page_number=i + 1, text=text))
        return pages

    def _extract_text(self, path: str) -> List[PageContent]:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        return [PageContent(page_number=1, text=text)]

    def _extract_docx(self, path: str) -> List[PageContent]:
        from docx import Document
        doc = Document(path)
        # Group paragraphs into ~50-paragraph "pages" to keep chunks manageable
        all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        page_size = 50
        pages = []
        for i in range(0, max(len(all_paras), 1), page_size):
            chunk = all_paras[i:i + page_size]
            pages.append(PageContent(
                page_number=len(pages) + 1,
                text="\n".join(chunk),
            ))
        return pages or [PageContent(page_number=1, text="")]

    def _is_scanned(self, pages: List[PageContent]) -> bool:
        avg = sum(len(p.text) for p in pages) / max(len(pages), 1)
        return avg < 50
